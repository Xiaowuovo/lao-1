"""
增强文件处理器 - 支持PDF、Word、图片OCR
"""

import os
import io
from werkzeug.utils import secure_filename


class EnhancedFileProcessor:
    """增强文件处理器"""
    
    def __init__(self):
        self.supported_extensions = {
            'text': ['.txt'],
            'pdf': ['.pdf'],
            'word': ['.doc', '.docx'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
        }
    
    def is_allowed_file(self, filename):
        """检查文件是否允许"""
        if not filename:
            return False
        
        ext = os.path.splitext(filename)[1].lower()
        for file_types in self.supported_extensions.values():
            if ext in file_types:
                return True
        return False
    
    def get_file_type(self, filename):
        """获取文件类型"""
        ext = os.path.splitext(filename)[1].lower()
        
        for file_type, extensions in self.supported_extensions.items():
            if ext in extensions:
                return file_type
        
        return 'unknown'
    
    def process_file(self, file, file_type=None):
        """
        处理文件并提取文本
        
        Args:
            file: 上传的文件对象
            file_type: 文件类型 (text/pdf/word/image)
            
        Returns:
            str: 提取的文本内容
        """
        if not file_type:
            file_type = self.get_file_type(file.filename)
        
        if file_type == 'text':
            return self.process_text_file(file)
        elif file_type == 'pdf':
            return self.process_pdf_file(file)
        elif file_type == 'word':
            return self.process_word_file(file)
        elif file_type == 'image':
            return self.process_image_file(file)
        else:
            raise ValueError(f'不支持的文件类型: {file_type}')
    
    def process_text_file(self, file):
        """处理文本文件"""
        try:
            # 尝试UTF-8编码
            content = file.read().decode('utf-8')
            return content
        except UnicodeDecodeError:
            # 尝试GBK编码
            file.seek(0)
            try:
                content = file.read().decode('gbk')
                return content
            except:
                raise ValueError('无法解码文本文件，请确保使用UTF-8或GBK编码')
    
    def process_pdf_file(self, file):
        """处理PDF文件"""
        try:
            import PyPDF2
            
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            
            if not text.strip():
                raise ValueError('PDF文件中没有可提取的文本内容')
            
            return text
        except ImportError:
            raise ValueError('PDF处理功能未安装，请安装：pip install PyPDF2')
        except Exception as e:
            raise ValueError(f'PDF文件处理失败: {str(e)}')
    
    def process_word_file(self, file):
        """处理Word文件"""
        try:
            import docx
            
            doc = docx.Document(file)
            text = ''
            
            # 提取段落
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + '\t'
                    text += '\n'
            
            if not text.strip():
                raise ValueError('Word文件中没有可提取的文本内容')
            
            return text
        except ImportError:
            raise ValueError('Word处理功能未安装，请安装：pip install python-docx')
        except Exception as e:
            raise ValueError(f'Word文件处理失败: {str(e)}')
    
    def process_image_file(self, file):
        """处理图片文件 - OCR识别"""
        try:
            from PIL import Image
            import pytesseract
            
            # 打开图片
            image = Image.open(file)
            
            # 图片预处理：转换为灰度图，增强对比度
            image = image.convert('L')  # 转为灰度
            
            # 配置OCR参数
            custom_config = r'--oem 3 --psm 6'
            
            # 进行OCR识别
            # 优先使用简体中文，然后尝试中英文混合
            try:
                # 首先尝试简体中文
                text = pytesseract.image_to_string(image, lang='chi_sim', config=custom_config)
            except:
                # 如果失败，尝试中英文混合
                try:
                    text = pytesseract.image_to_string(image, lang='chi_sim+eng', config=custom_config)
                except:
                    # 最后尝试仅英文
                    text = pytesseract.image_to_string(image, lang='eng', config=custom_config)
            
            if not text.strip():
                raise ValueError('图片中没有识别到文字内容，请确保图片清晰且包含可识别的文字')
            
            return text
        except ImportError as e:
            if 'PIL' in str(e) or 'Image' in str(e):
                raise ValueError('图片处理功能未安装，请安装：pip install Pillow')
            elif 'pytesseract' in str(e):
                raise ValueError('OCR功能未安装，请安装：pip install pytesseract 并安装Tesseract-OCR')
            else:
                raise ValueError(f'导入模块失败: {str(e)}')
        except Exception as e:
            raise ValueError(f'图片OCR识别失败: {str(e)}')


# 单例
_processor_instance = None

def get_file_processor():
    """获取文件处理器单例"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = EnhancedFileProcessor()
    return _processor_instance


if __name__ == '__main__':
    # 测试代码
    processor = EnhancedFileProcessor()
    
    # 测试文件类型识别
    test_files = [
        'test.txt',
        'document.pdf',
        'report.docx',
        'image.jpg',
        'unknown.xyz'
    ]
    
    print("=" * 60)
    print("文件类型识别测试")
    print("=" * 60)
    
    for filename in test_files:
        file_type = processor.get_file_type(filename)
        is_allowed = processor.is_allowed_file(filename)
        print(f"{filename:20} -> 类型: {file_type:10} 允许: {is_allowed}")
