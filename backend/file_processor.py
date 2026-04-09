import os
import requests
from typing import Dict, Optional
from bs4 import BeautifulSoup
import pymysql

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class FileProcessor:
    """多格式文件处理器"""
    
    def __init__(self, db_config: dict, upload_folder: str = 'uploads'):
        self.db_config = db_config
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)
    
    def process_image(self, image_path: str, user_id: int) -> Dict:
        """
        处理图片，提取文字（OCR）
        
        Returns:
            {
                'success': True/False,
                'text': '提取的文字',
                'confidence': 0.95,
                'error': ''
            }
        """
        if not HAS_OCR:
            return {
                'success': False,
                'text': '',
                'confidence': 0.0,
                'error': 'OCR功能未安装，请安装: pip install pytesseract pillow'
            }
        
        try:
            # 打开图片
            img = Image.open(image_path)
            
            # 执行OCR识别
            # 配置中文识别
            config = '--psm 6 -l chi_sim+eng'
            text = pytesseract.image_to_string(img, config=config)
            
            # 计算置信度（简化版）
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, config=config)
            confidences = [int(conf) for conf in data['conf'] if conf != '-1']
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # 记录到数据库
            file_size = os.path.getsize(image_path)
            self._record_file(user_id, os.path.basename(image_path), 'image', 
                            image_path, file_size, text, avg_confidence / 100)
            
            return {
                'success': True,
                'text': text.strip(),
                'confidence': round(avg_confidence / 100, 2),
                'error': ''
            }
        
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'confidence': 0.0,
                'error': f'OCR识别失败: {str(e)}'
            }
    
    def process_pdf(self, pdf_path: str, user_id: int) -> Dict:
        """
        处理PDF文件，提取文字
        
        Returns:
            {
                'success': True/False,
                'text': '提取的文字',
                'pages': 5,
                'error': ''
            }
        """
        if not HAS_PDF:
            return {
                'success': False,
                'text': '',
                'pages': 0,
                'error': 'PDF处理功能未安装，请安装: pip install PyPDF2'
            }
        
        try:
            text = ''
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pages = len(reader.pages)
                
                # 提取所有页面文字
                for page in reader.pages:
                    text += page.extract_text() + '\n'
            
            # 记录到数据库
            file_size = os.path.getsize(pdf_path)
            self._record_file(user_id, os.path.basename(pdf_path), 'pdf', 
                            pdf_path, file_size, text, 1.0)
            
            return {
                'success': True,
                'text': text.strip(),
                'pages': pages,
                'error': ''
            }
        
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'pages': 0,
                'error': f'PDF处理失败: {str(e)}'
            }
    
    def process_word(self, word_path: str, user_id: int) -> Dict:
        """
        处理Word文档，提取文字
        
        Returns:
            {
                'success': True/False,
                'text': '提取的文字',
                'paragraphs': 10,
                'error': ''
            }
        """
        if not HAS_DOCX:
            return {
                'success': False,
                'text': '',
                'paragraphs': 0,
                'error': 'Word处理功能未安装，请安装: pip install python-docx'
            }
        
        try:
            doc = docx.Document(word_path)
            
            # 提取所有段落
            text = '\n'.join([para.text for para in doc.paragraphs])
            paragraphs = len(doc.paragraphs)
            
            # 记录到数据库
            file_size = os.path.getsize(word_path)
            self._record_file(user_id, os.path.basename(word_path), 'word', 
                            word_path, file_size, text, 1.0)
            
            return {
                'success': True,
                'text': text.strip(),
                'paragraphs': paragraphs,
                'error': ''
            }
        
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'paragraphs': 0,
                'error': f'Word处理失败: {str(e)}'
            }
    
    def process_url(self, url: str, user_id: int) -> Dict:
        """
        处理网页链接，爬取文字内容
        
        Returns:
            {
                'success': True/False,
                'text': '爬取的文字',
                'title': '网页标题',
                'url': 'http://...',
                'error': ''
            }
        """
        try:
            # 发送请求
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.encoding = response.apparent_encoding
            
            # 解析HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 提取标题
            title = soup.title.string if soup.title else '未知标题'
            
            # 移除脚本和样式
            for script in soup(['script', 'style']):
                script.decompose()
            
            # 提取正文（尝试找到主要内容区域）
            main_content = soup.find('article') or soup.find('main') or soup.find('div', class_=['content', 'main', 'article'])
            
            if main_content:
                text = main_content.get_text(separator='\n', strip=True)
            else:
                text = soup.get_text(separator='\n', strip=True)
            
            # 清理多余空行
            text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
            
            # 记录到数据库
            self._record_file(user_id, url, 'url', url, 0, text, 1.0)
            
            return {
                'success': True,
                'text': text[:5000],  # 限制长度
                'title': title,
                'url': url,
                'error': ''
            }
        
        except requests.RequestException as e:
            return {
                'success': False,
                'text': '',
                'title': '',
                'url': url,
                'error': f'网页爬取失败: {str(e)}'
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'title': '',
                'url': url,
                'error': f'网页解析失败: {str(e)}'
            }
    
    def _record_file(self, user_id: int, file_name: str, file_type: str,
                    file_path: str, file_size: int, extracted_text: str, 
                    confidence: float):
        """记录文件处理信息到数据库"""
        try:
            conn = pymysql.connect(**self.db_config)
            cursor = conn.cursor()
            
            cursor.execute("""
                INSERT INTO uploaded_files 
                (user_id, file_name, file_type, file_path, file_size, 
                 extracted_text, ocr_confidence, processing_status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, 'completed')
            """, (user_id, file_name, file_type, file_path, file_size, 
                  extracted_text, confidence))
            
            conn.commit()
            cursor.close()
            conn.close()
        except Exception as e:
            print(f"记录文件信息失败: {e}")
    
    def get_file_history(self, user_id: int, limit: int = 20) -> list:
        """获取用户的文件上传历史"""
        try:
            conn = pymysql.connect(**self.db_config)
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            
            cursor.execute("""
                SELECT * FROM uploaded_files
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT %s
            """, (user_id, limit))
            
            files = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            return files
        except Exception as e:
            print(f"获取文件历史失败: {e}")
            return []


class BatchTextSplitter:
    """批量文本拆分器"""
    
    @staticmethod
    def split_multiple_events(text: str) -> list:
        """
        将包含多条通知的文本拆分为多条
        
        策略：
        1. 按空行分割
        2. 按编号分割（1. 2. 3.）
        3. 按日期分割
        """
        events = []
        
        # 方法1：按空行分割
        parts = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if len(parts) > 1:
            return parts
        
        # 方法2：按序号分割
        import re
        numbered_pattern = r'(?:^|\n)(\d+[.、]\s*.+?)(?=\n\d+[.、]|\Z)'
        matches = re.findall(numbered_pattern, text, re.DOTALL)
        
        if matches and len(matches) > 1:
            return [m.strip() for m in matches]
        
        # 方法3：按日期分割（保留日期）
        date_pattern = r'(\d{1,2}月\d{1,2}日[^\n]+(?:\n[^\d月\n]+)*)'
        date_matches = re.findall(date_pattern, text)
        
        if date_matches and len(date_matches) > 1:
            return [m.strip() for m in date_matches]
        
        # 如果无法拆分，返回整段文本
        return [text]
    
    @staticmethod
    def clean_text(text: str) -> str:
        """清理文本"""
        # 移除多余空格和换行
        text = re.sub(r'\s+', ' ', text)
        # 移除特殊字符
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。！？、；：""''（）《》【】\-/:：]', '', text)
        return text.strip()
